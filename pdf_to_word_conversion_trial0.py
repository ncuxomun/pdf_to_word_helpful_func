import os
import camelot
import docx
from PIL import Image

def pdf_to_word(pdf_file, word_file):
    """Convert PDF to Word doc while maintaining structure"""
    
    # Load the PDF
    tables = camelot.read_pdf(pdf_file)
    
    # Create Word doc
    doc = docx.Document()
    
    # Add tables
    for table in tables:
        doc_table = doc.add_table(rows=table.df.shape[0], cols=table.df.shape[1])
        for i, row in enumerate(table.df.values):
            for j, val in enumerate(row):
                doc_table.cell(i,j).text = str(val)
                
    # Add images
    for i, img in enumerate(images):
        doc.add_picture(img, width=docx.shared.Inches(6))
        
    # Save Word doc
    doc.save(word_file)
# Example usage
    
file_path = os.getcwd() + r"\xxxx.pdf"
output_path = os.getcwd() + r"\xxx.docx"

pdf_to_word(file_path, output_path)
